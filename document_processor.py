#!/usr/bin/env python3
"""
Document Processor for extracting and analyzing email attachments
"""

import io
import re
import numpy as np
from typing import Dict, List, Any, Optional
from pathlib import Path
import tempfile
import os
import pandas as pd

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class DocumentProcessor:
    """Process and extract text from various document types"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/vnd.ms-excel': self._extract_excel_text,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_excel_text,
            'text/plain': self._extract_text_text,
            'text/html': self._extract_html_text,
            'text/csv': self._extract_csv_text
        }
    
    def extract_document_text(self, attachment_data: bytes, mime_type: str, filename: str) -> Dict[str, Any]:
        """Extract text from document attachment"""
        try:
            if mime_type in self.supported_types:
                extractor = self.supported_types[mime_type]
                text_content = extractor(attachment_data, filename)
                
                return {
                    'success': True,
                    'text': text_content,
                    'word_count': len(text_content.split()),
                    'filename': filename,
                    'mime_type': mime_type,
                    'extracted': True
                }
            else:
                return {
                    'success': False,
                    'text': '',
                    'word_count': 0,
                    'filename': filename,
                    'mime_type': mime_type,
                    'extracted': False,
                    'error': f'Unsupported file type: {mime_type}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'word_count': 0,
                'filename': filename,
                'mime_type': mime_type,
                'extracted': False,
                'error': str(e)
            }
    
    def _extract_pdf_text(self, data: bytes, filename: str) -> str:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            return f"[PDF file: {filename} - PyPDF2 not available for text extraction]"
        
        try:
            pdf_file = io.BytesIO(data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            return f"[PDF extraction error: {str(e)}]"
    
    def _extract_docx_text(self, data: bytes, filename: str) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            return f"[DOCX file: {filename} - python-docx not available for text extraction]"
        
        try:
            doc = Document(io.BytesIO(data))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            return f"[DOCX extraction error: {str(e)}]"
    
    def _extract_excel_text(self, data: bytes, filename: str) -> str:
        """Extract text from Excel files with intelligent sheet handling"""
        if not EXCEL_AVAILABLE:
            return f"[Excel file: {filename} - pandas not available for text extraction]"
        
        try:
            excel_file = io.BytesIO(data)
            
            # Try to read as Excel
            try:
                df_dict = pd.read_excel(excel_file, sheet_name=None)
            except:
                # Try as CSV
                excel_file.seek(0)
                df = pd.read_csv(excel_file)
                df_dict = {'Sheet1': df}
            
            text = ""
            sheet_analysis = []
            
            for sheet_name, sheet_df in df_dict.items():
                # Analyze each sheet
                sheet_info = self._analyze_excel_sheet(sheet_name, sheet_df)
                sheet_analysis.append(sheet_info)
                
                # Add sheet summary to text
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += f"Rows: {sheet_info['rows']}, Columns: {sheet_info['columns']}\n"
                
                if sheet_info['has_data']:
                    # Add column names if they exist
                    if sheet_info['column_names']:
                        text += f"Columns: {', '.join(sheet_info['column_names'][:10])}"
                        if len(sheet_info['column_names']) > 10:
                            text += f" ... and {len(sheet_info['column_names']) - 10} more"
                        text += "\n"
                    
                    # Add sample data (first few rows)
                    if sheet_info['sample_data']:
                        text += "Sample data:\n"
                        text += sheet_info['sample_data']
                        text += "\n"
                    
                    # Add key statistics if available
                    if sheet_info['statistics']:
                        text += f"Key statistics: {sheet_info['statistics']}\n"
                else:
                    text += "Sheet appears to be empty or contains minimal data\n"
            
            # Add overall analysis
            text += f"\n--- Overall Analysis ---\n"
            text += f"Total sheets: {len(df_dict)}\n"
            text += f"Sheets with data: {sum(1 for s in sheet_analysis if s['has_data'])}\n"
            
            # Identify the most important sheet
            if sheet_analysis:
                main_sheet = max(sheet_analysis, key=lambda x: x['rows'] * x['columns'])
                text += f"Main sheet: {main_sheet['name']} ({main_sheet['rows']} rows, {main_sheet['columns']} columns)\n"
            
            return text.strip()
            
        except Exception as e:
            return f"[Excel extraction error: {str(e)}]"
    
    def _analyze_excel_sheet(self, sheet_name: str, df: pd.DataFrame) -> dict:
        """Analyze a single Excel sheet and extract meaningful information"""
        try:
            rows, cols = df.shape
            
            # Check if sheet has meaningful data
            has_data = rows > 0 and cols > 0
            
            # Get column names (handle unnamed columns)
            column_names = []
            if has_data:
                for i, col in enumerate(df.columns):
                    if pd.isna(col) or str(col).startswith('Unnamed'):
                        column_names.append(f"Column_{i+1}")
                    else:
                        column_names.append(str(col))
            
            # Get sample data (first 3 rows, limited columns)
            sample_data = ""
            if has_data and rows > 0:
                # Take first 3 rows and first 5 columns
                sample_df = df.head(3).iloc[:, :5]
                sample_data = sample_df.to_string(max_cols=5, max_rows=3)
            
            # Extract key statistics for numeric columns
            statistics = []
            if has_data:
                numeric_cols = df.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 0:
                    for col in numeric_cols[:3]:  # Limit to first 3 numeric columns
                        col_stats = df[col].describe()
                        if not col_stats.isna().all():
                            mean_val = col_stats['mean']
                            if not pd.isna(mean_val):
                                statistics.append(f"{col}: avg={mean_val:.2f}")
            
            return {
                'name': sheet_name,
                'rows': rows,
                'columns': cols,
                'has_data': has_data,
                'column_names': column_names,
                'sample_data': sample_data,
                'statistics': ', '.join(statistics) if statistics else None
            }
            
        except Exception as e:
            return {
                'name': sheet_name,
                'rows': 0,
                'columns': 0,
                'has_data': False,
                'column_names': [],
                'sample_data': f"Error analyzing sheet: {str(e)}",
                'statistics': None
            }
    
    def _extract_text_text(self, data: bytes, filename: str) -> str:
        """Extract text from plain text files"""
        try:
            return data.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"[Text extraction error: {str(e)}]"
    
    def _extract_html_text(self, data: bytes, filename: str) -> str:
        """Extract text from HTML files"""
        try:
            html_content = data.decode('utf-8', errors='ignore')
            # Simple HTML tag removal
            text = re.sub(r'<[^>]+>', '', html_content)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
        except Exception as e:
            return f"[HTML extraction error: {str(e)}]"
    
    def _extract_csv_text(self, data: bytes, filename: str) -> str:
        """Extract text from CSV files"""
        if not EXCEL_AVAILABLE:
            return f"[CSV file: {filename} - pandas not available for text extraction]"
        
        try:
            csv_file = io.BytesIO(data)
            df = pd.read_csv(csv_file)
            return df.to_string()
        except Exception as e:
            return f"[CSV extraction error: {str(e)}]"
    
    def analyze_document_content(self, text: str, filename: str) -> Dict[str, Any]:
        """Analyze extracted document content"""
        if not text or len(text.strip()) < 10:
            return {
                'success': False,
                'analysis': 'Document appears to be empty or contains minimal text',
                'key_points': [],
                'document_type': 'unknown'
            }
        
        # Basic content analysis
        lines = text.split('\n')
        words = text.split()
        
        # Detect document type based on content
        doc_type = self._detect_document_type(text, filename)
        
        # Extract key points (improved implementation)
        key_points = self._extract_key_points(text)
        
        return {
            'success': True,
            'analysis': f"Document contains {len(words)} words across {len(lines)} lines",
            'key_points': key_points,
            'document_type': doc_type,
            'word_count': len(words),
            'line_count': len(lines)
        }
    
    def _detect_document_type(self, text: str, filename: str) -> str:
        """Detect document type based on content and filename"""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Excel/Spreadsheet detection (check first)
        if any(keyword in filename_lower for keyword in ['.xlsx', '.xls', '.csv']):
            # Check if it's a financial report
            if any(keyword in text_lower for keyword in ['transaction', 'amount', 'balance', 'total', 'summary', 'cdf']):
                return 'financial_report'
            return 'spreadsheet'
        
        # Contract/Agreement detection
        if any(keyword in text_lower for keyword in ['contract', 'agreement', 'terms', 'conditions', 'party', 'signature']):
            return 'contract'
        
        # Report detection
        if any(keyword in text_lower for keyword in ['report', 'summary', 'conclusion', 'findings', 'analysis']):
            return 'report'
        
        # Invoice/Billing detection
        if any(keyword in text_lower for keyword in ['invoice', 'bill', 'payment', 'amount', 'total', 'due']):
            return 'invoice'
        
        # Resume/CV detection
        if any(keyword in text_lower for keyword in ['resume', 'cv', 'experience', 'education', 'skills']):
            return 'resume'
        
        # Proposal detection
        if any(keyword in text_lower for keyword in ['proposal', 'project', 'scope', 'deliverables', 'timeline']):
            return 'proposal'
        
        return 'document'
    
    def _extract_key_points(self, text: str) -> List[str]:
        """Extract key points from document text"""
        key_points = []
        
        # Special handling for Excel/spreadsheet documents
        if '--- Sheet:' in text and ('Rows:' in text or 'Columns:' in text):
            # Extract sheet information
            sheet_sections = text.split('--- Sheet:')
            for section in sheet_sections[1:]:  # Skip the first empty section
                lines = section.split('\n')
                if len(lines) > 1:
                    sheet_name = lines[0].strip()
                    for line in lines[1:]:
                        if 'Rows:' in line and 'Columns:' in line:
                            key_points.append(f"Sheet '{sheet_name}': {line.strip()}")
                        elif 'Main sheet:' in line:
                            key_points.append(f"Main data sheet: {line.strip()}")
                        elif 'Key statistics:' in line and line.strip() != 'Key statistics:':
                            key_points.append(f"Data insights: {line.strip()}")
                        elif 'Total sheets:' in line:
                            key_points.append(f"Document structure: {line.strip()}")
            
            # If no specific points found, add general structure info
            if not key_points:
                if 'Total sheets:' in text:
                    total_sheets_match = re.search(r'Total sheets: (\d+)', text)
                    if total_sheets_match:
                        key_points.append(f"Multi-sheet document with {total_sheets_match.group(1)} sheets")
                
                if 'Main sheet:' in text:
                    main_sheet_match = re.search(r'Main sheet: ([^(]+)', text)
                    if main_sheet_match:
                        key_points.append(f"Primary data in sheet: {main_sheet_match.group(1).strip()}")
        
        # General text analysis for non-Excel documents
        else:
            # Split into sentences
            sentences = re.split(r'[.!?]+', text)
            
            # Look for sentences with key indicators
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) < 10:
                    continue
                    
                # Look for sentences with numbers, dates, or key terms
                if (re.search(r'\d+', sentence) or 
                    any(keyword in sentence.lower() for keyword in ['important', 'key', 'critical', 'urgent', 'deadline', 'due'])):
                    key_points.append(sentence[:200] + '...' if len(sentence) > 200 else sentence)
        
        # Limit to top 5 key points
        return key_points[:5] 